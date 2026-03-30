from django.http import JsonResponse
from .models import Lab, LabMembership, Activitate, UserProfile, WorkEntry, MonthlyMeta
from django.views.decorators.csrf import csrf_exempt
import json
from django.contrib.auth.models import User
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_http_methods
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404, render, redirect
from django.utils.http import url_has_allowed_host_and_scheme
from .serializers import WorkEntrySerializer
import openpyxl
from datetime import datetime
import requests
from django.conf import settings
from rest_framework.decorators import api_view
from rest_framework.response import Response
from django.db.models import Sum
from django.contrib.auth import authenticate, login, logout
from django.http import HttpResponseForbidden
import calendar
from collections import defaultdict
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Alignment, Font
from openpyxl.styles.borders import Border, Side
from io import BytesIO
from . import export_views

LAB_JURNAL_NAMES = ["Lab1", "Lab2"]

def _is_admin(user):
    return getattr(getattr(user, "userprofile", None), "role", None) == "admin"

def _is_director_for_lab(user, lab_id):
    if not lab_id:
        return False
    return LabMembership.objects.filter(
        profile=user.userprofile,
        lab_id=lab_id,
        role="director",
    ).exists()

@csrf_exempt
def login_page(request):
    next_url = request.GET.get("next") or request.POST.get("next") or ""
    if request.method == "POST":
        username = request.POST.get("username")
        password = request.POST.get("password")

        user = authenticate(request, username=username, password=password)

        if user is not None:
            login(request, user)
            if next_url and url_has_allowed_host_and_scheme(
                next_url, allowed_hosts={request.get_host()}, require_https=request.is_secure()
            ):
                return redirect(next_url)
            return redirect("index")  # redirect to your main page
        else:
            return render(
                request, "api/login.html", {"error": "Invalid credentials", "next": next_url}
            )

    return render(request, "api/login.html", {"next": next_url})
    
@api_view(['POST'])
def logout_view(request):
    logout(request)
    return JsonResponse({"message": "Logged out"})

@login_required
def get_pontaj_dates(request):
    '''Endpoint to get all dates where the user has work entries, for a given lab.'''
    lab_id = request.GET.get("lab")

    entries = WorkEntry.objects.filter(
        user=request.user,
        lab_id=lab_id
    )

    dates = list(entries.values_list("date", flat=True))

    formatted_dates = [d.strftime("%Y-%m-%d") for d in dates]

    return JsonResponse(formatted_dates, safe=False)

@login_required
def get_monthly_hours(request):
    '''Endpoint to get total hours worked by the user in a given month/year and lab, along with the monthly limit.'''
    user = request.user
    # read month/year from query params, default to current
    month = int(request.GET.get("month", datetime.now().month))
    year = int(request.GET.get("year", datetime.now().year))

    lab_id = request.GET.get("lab")
    if not lab_id:
        return JsonResponse({"error": "lab_id required"}, status=400)

    membership = LabMembership.objects.filter(
        profile=request.user.userprofile,
        lab_id=lab_id
    ).first()

    if not membership:
        return JsonResponse({
            "used_hours": 0,
            "limit": 0, 
            "remaining": 0
        })

    total = WorkEntry.objects.filter(
        user=request.user,
        lab_id=lab_id,
        date__month=month,
        date__year=year
    ).aggregate(total=Sum("nr_ore"))["total"] or 0

    limit = membership.monthly_hour_limit

    return JsonResponse({
        "used_hours": total,
        "limit": limit,
        "remaining": limit - total
    })

# @login_required
# def current_user(request):
#     if request.user.is_authenticated:
#         return JsonResponse({
#             "username": request.user.username,
#             "role": request.user.userprofile.role
#         })
#     return JsonResponse({"username": None})

# @login_required
# def current_user(request):
#     '''Endpoint to get current user info, including global role and lab-specific role if lab_id is provided.'''
#     user = request.user
#     profile = user.userprofile

#     lab_id = request.GET.get("lab_id")

#     lab_role = None
#     is_director = False

#     if lab_id:
#         membership = LabMembership.objects.filter(
#             profile=profile,
#             lab_id=lab_id
#         ).first()

#         if membership:
#             lab_role = membership.role
#             is_director = membership.role == "director"

#     data = {
#         "username": user.username,
#         "global_role": profile.role,  # admin or not
#         "lab_role": lab_role,
#         "is_director": is_director,
#     }

#     return JsonResponse(data)

@login_required
def current_user(request):
    profile = request.user.userprofile
    lab_id = request.GET.get("lab")

    lab_role = None
    can_see_jurnal = LabMembership.objects.filter(
        profile=profile,
        lab__name__in=LAB_JURNAL_NAMES
    ).exists()
    print("CAN SEE JURNAL:", can_see_jurnal)

    if lab_id:
        try:
            membership = LabMembership.objects.get(
                lab_id=lab_id,
                profile=profile
            )
            lab_role = membership.role
        except LabMembership.DoesNotExist:
            pass

    return JsonResponse({
        "username": request.user.username,
        "global_role": profile.role,
        "lab_role": lab_role,
        "can_see_jurnal": can_see_jurnal,
    })


@login_required
def index(request):
    profile = request.user.userprofile
    can_see_jurnal = LabMembership.objects.filter(
        profile=profile,
        lab__name__in=LAB_JURNAL_NAMES
    ).exists()
    return render(request, "api/index.html", {"can_see_jurnal": can_see_jurnal})

@login_required
def entries_page(request):
    profile = request.user.userprofile
    can_see_jurnal = LabMembership.objects.filter(
        profile=profile,
        lab__name__in=LAB_JURNAL_NAMES
    ).exists()
    return render(request, "api/entries.html", {"can_see_jurnal": can_see_jurnal})

@login_required
def members_hours_page(request):
    """
    Page for directors/admins to see every member's worked hours (nr_ore) per day in a month.
    """
    profile = request.user.userprofile
    is_admin = _is_admin(request.user)
    is_director_any = LabMembership.objects.filter(profile=profile, role="director").exists()

    if not (is_admin or is_director_any):
        return HttpResponseForbidden("Nu ai acces la aceasta pagina.")

    can_see_jurnal = LabMembership.objects.filter(
        profile=profile,
        lab__name__in=LAB_JURNAL_NAMES,
    ).exists()

    return render(request, "api/members_hours.html", {"can_see_jurnal": can_see_jurnal})

@login_required
def annual_stats_page(request):
    """
    Page for directors/admins to see yearly statistics:
    per member totals for each month + year total (nr_ore only).
    """
    profile = request.user.userprofile
    is_admin = _is_admin(request.user)
    is_director_any = LabMembership.objects.filter(profile=profile, role="director").exists()

    if not (is_admin or is_director_any):
        return HttpResponseForbidden("Nu ai acces la aceasta pagina.")

    can_see_jurnal = LabMembership.objects.filter(
        profile=profile,
        lab__name__in=LAB_JURNAL_NAMES,
    ).exists()

    return render(request, "api/annual_stats.html", {"can_see_jurnal": can_see_jurnal})

@login_required
def list_director_labs(request):
    """
    Labs that the user can manage (director in that lab) or all labs for global admin.
    Used by the members-hours page.
    """
    profile = request.user.userprofile

    if _is_admin(request.user):
        labs = Lab.objects.all().order_by("id")
    else:
        labs = (
            Lab.objects.filter(labmembership__profile=profile, labmembership__role="director")
            .distinct()
            .order_by("id")
        )

    data = [{"id": l.id, "name": l.name} for l in labs]
    return JsonResponse(data, safe=False)

@login_required
def get_members_monthly_hours(request):
    """
    Data endpoint for the members-hours page.

    Returns worked hours (nr_ore) per user per day for the specified lab/month/year.
    Visibility: global admin OR director of that lab.
    """
    lab_id = request.GET.get("lab")
    month = request.GET.get("month")
    year = request.GET.get("year")

    try:
        lab_id_int = int(lab_id)
        month_int = int(month)
        year_int = int(year)
    except (TypeError, ValueError):
        return JsonResponse({"error": "Invalid lab/month/year"}, status=400)

    if not (_is_admin(request.user) or _is_director_for_lab(request.user, lab_id_int)):
        return HttpResponseForbidden("Nu ai acces la acest lab.")

    if month_int < 1 or month_int > 12:
        return JsonResponse({"error": "Invalid month"}, status=400)

    days_in_month = calendar.monthrange(year_int, month_int)[1]
    lab = get_object_or_404(Lab, id=lab_id_int)

    memberships = (
        LabMembership.objects.filter(lab=lab)
        .select_related("profile__user")
    )
    users = [m.profile.user for m in memberships]

    users_sorted = sorted(
        users,
        key=lambda u: (
            (u.last_name or "").lower(),
            (u.first_name or "").lower(),
            (u.username or "").lower(),
        ),
    )

    # Aggregate hours per (user, date).
    aggregates = (
        WorkEntry.objects.filter(
            lab=lab,
            user__in=users_sorted,
            date__year=year_int,
            date__month=month_int,
        )
        .values("user_id", "date")
        .annotate(hours=Sum("nr_ore"))
    )

    hours_by_user_day = defaultdict(lambda: defaultdict(int))
    for row in aggregates:
        date_obj = row["date"]
        if not date_obj:
            continue
        day = int(date_obj.day)
        hours_by_user_day[int(row["user_id"])][day] = int(row["hours"] or 0)

    members = []
    for u in users_sorted:
        daily = [0] * days_in_month
        total = 0
        by_day = hours_by_user_day.get(u.id, {})
        for d in range(1, days_in_month + 1):
            h = int(by_day.get(d, 0) or 0)
            daily[d - 1] = h
            total += h

        display_name = f"{(u.last_name or '').strip()} {(u.first_name or '').strip()}".strip()
        if not display_name:
            display_name = u.username

        members.append(
            {
                "user_id": u.id,
                "name": display_name,
                "daily_hours": daily,
                "total": total,
            }
        )

    return JsonResponse(
        {
            "lab": {"id": lab.id, "name": lab.name},
            "month": month_int,
            "year": year_int,
            "days_in_month": days_in_month,
            "members": members,
        }
    )

@login_required
def get_members_yearly_hours(request):
    """
    Data endpoint for the annual-stats page.

    Returns worked hours (nr_ore) per member per month for the specified lab/year.
    Visibility: global admin OR director of that lab.
    """
    lab_id = request.GET.get("lab")
    year = request.GET.get("year")

    try:
        lab_id_int = int(lab_id)
        year_int = int(year)
    except (TypeError, ValueError):
        return JsonResponse({"error": "Invalid lab/year"}, status=400)

    if not (_is_admin(request.user) or _is_director_for_lab(request.user, lab_id_int)):
        return HttpResponseForbidden("Nu ai acces la acest lab.")

    lab = get_object_or_404(Lab, id=lab_id_int)

    memberships = (
        LabMembership.objects.filter(lab=lab)
        .select_related("profile__user")
    )
    users = [m.profile.user for m in memberships]

    users_sorted = sorted(
        users,
        key=lambda u: (
            (u.last_name or "").lower(),
            (u.first_name or "").lower(),
            (u.username or "").lower(),
        ),
    )

    aggregates = (
        WorkEntry.objects.filter(
            lab=lab,
            user__in=users_sorted,
            date__year=year_int,
        )
        .values("user_id", "date__month")
        .annotate(hours=Sum("nr_ore"))
    )

    hours_by_user_month = defaultdict(lambda: defaultdict(int))
    for row in aggregates:
        month_val = row.get("date__month")
        if not month_val:
            continue
        hours_by_user_month[int(row["user_id"])][int(month_val)] = int(row["hours"] or 0)

    members = []
    for u in users_sorted:
        monthly = [0] * 12
        total = 0
        by_month = hours_by_user_month.get(u.id, {})
        for m in range(1, 13):
            h = int(by_month.get(m, 0) or 0)
            monthly[m - 1] = h
            total += h

        display_name = f"{(u.last_name or '').strip()} {(u.first_name or '').strip()}".strip()
        if not display_name:
            display_name = u.username

        members.append(
            {
                "user_id": u.id,
                "name": display_name,
                "monthly_hours": monthly,
                "total": total,
            }
        )

    return JsonResponse(
        {
            "lab": {"id": lab.id, "name": lab.name},
            "year": year_int,
            "members": members,
        }
    )

def get_visible_labs(user):
    '''Helper function to get labs visible to the user based on their role. Admin sees all, director sees their labs, member sees their labs.'''
    profile = user.userprofile

    if profile.role == "admin":
        return Lab.objects.all()

    if profile.role == "director":
        return profile.labs.all()

    return profile.labs.all()

@login_required
def list_labs(request):
    labs = get_visible_labs(request.user)
    data = [{"id": l.id, "name": l.name} for l in labs]
    return JsonResponse(data, safe=False)

# def list_labs(request):
#     labs = Lab.objects.all()
#     data = [{"id": l.id, "name": l.name} for l in labs]
#     return JsonResponse(data, safe=False)


def list_activitati(request, lab_id):
    subs = Activitate.objects.filter(lab_id=lab_id)
    data = [{
        "id": s.id,
        "nume": s.nume,
        "descriere": s.descriere,
    } for s in subs]
    return JsonResponse(data, safe=False)


@require_http_methods(["POST"])
def create_work_entry(request):
    '''Endpoint to create a work entry. Validates that the user is enrolled in the lab and does not exceed monthly hour limit.'''
    if request.method == "POST":
        data = json.loads(request.body)
        print("RECEIVED:", data)

        user = request.user
        profile = user.userprofile
        lab_id = data.get("lab")

        can_see_jurnal = LabMembership.objects.filter(
            profile=profile,
            lab__name__in=LAB_JURNAL_NAMES
        ).exists()
        if not can_see_jurnal:
            data.pop("jurnal", None)
            data.pop("scurta_descriere_jurnal", None)

        nr_ore = data.get("nr_ore")
        durata = data.get("durata")
        members = data.get("members", [])
        if not lab_id:
            return JsonResponse({"error": "Missing lab"}, status=400)

        if not data.get("nr_ore") or not data.get("durata"):
            return JsonResponse({"error": "Missing nr_ore or durata"}, status=400)
        
        if not data.get("nr_ore") or not data.get("durata"):
            return JsonResponse({"error": "Missing nr_ore or durata"}, status=400)
        
        membership = LabMembership.objects.filter(
            profile=profile,
            lab_id=lab_id
        ).first()

        if not membership and profile.role != "admin":
            return JsonResponse(
                {"error": "User not enrolled in this lab"},
                status=403
            )

        # Monthly fields (links/livrabil/comentarii) are now user-entered.

        date_obj = datetime.strptime(data["date"], "%Y-%m-%d")
        if date_obj.weekday() >= 5:  # 5=Saturday, 6=Sunday
            return JsonResponse(
                {"error": "Ați încercat să pontați in weekend"},
                status=400,
            )
        month = date_obj.month
        year = date_obj.year

        existing_hours = WorkEntry.objects.filter(
            user=user,
            lab_id=lab_id,
            date__month=month,
            date__year=year
        ).aggregate(total=Sum("nr_ore"))["total"] or 0

        try:
            new_hours = int(data["nr_ore"])
        except (TypeError, ValueError):
            return JsonResponse({"error": "nr_ore must be an integer"}, status=400)

        if new_hours < 1 or new_hours > 12:
            return JsonResponse({"error": "nr_ore must be between 1 and 12"}, status=400)

        # Directors/members get per-lab limits via LabMembership. Admins may not have
        # a LabMembership row, so fall back to their profile limit.
        limit = (
            membership.monthly_hour_limit
            if membership is not None
            else profile.monthly_hour_limit
        )
        remaining = limit - existing_hours

        if existing_hours + new_hours > limit:
            return JsonResponse(
                {"error": f"Ți-ai atins limita lunară. Mai poți introduce doar {remaining} ore."},
                status=400
            )

        serializer = WorkEntrySerializer(data=data)

        if serializer.is_valid():
            entry = serializer.save(user=user)

            if members:
                users = User.objects.filter(id__in=members)
                entry.members.set(users)
            return JsonResponse(serializer.data, status=201)

        print("SERIALIZER ERRORS:", serializer.errors)
        return JsonResponse(serializer.errors, status=400)

    return JsonResponse({"error": "POST only"}, status=400)


@login_required
@require_http_methods(["GET", "POST"])
def monthly_meta(request):
    """
    Upsert/read per-user monthly fields used by Conti export.
    Key: (user, lab, activitate, year, month)
    """
    profile = request.user.userprofile

    if request.method == "GET":
        lab_id = request.GET.get("lab")
        activitate_id = request.GET.get("activitate")
        month = request.GET.get("month")
        year = request.GET.get("year")
    else:
        try:
            payload = json.loads(request.body or "{}")
        except json.JSONDecodeError:
            return JsonResponse({"error": "Invalid JSON"}, status=400)

        lab_id = payload.get("lab")
        activitate_id = payload.get("activitate")
        month = payload.get("month")
        year = payload.get("year")

    try:
        lab_id_int = int(lab_id)
        activitate_id_int = int(activitate_id)
        month_int = int(month)
        year_int = int(year)
    except (TypeError, ValueError):
        return JsonResponse({"error": "Invalid lab/activitate/month/year"}, status=400)

    if month_int < 1 or month_int > 12:
        return JsonResponse({"error": "Invalid month"}, status=400)

    lab = get_object_or_404(Lab, id=lab_id_int)
    activitate = get_object_or_404(Activitate, id=activitate_id_int)

    if activitate.lab_id != lab.id:
        return JsonResponse({"error": "Activitate does not belong to selected Lab."}, status=400)

    is_admin = profile.role == "admin"
    if not is_admin:
        is_member = LabMembership.objects.filter(profile=profile, lab=lab).exists()
        if not is_member:
            return HttpResponseForbidden("User not enrolled in this lab")

    if request.method == "GET":
        meta = MonthlyMeta.objects.filter(
            user=request.user,
            lab=lab,
            activitate=activitate,
            month=month_int,
            year=year_int,
        ).first()
        return JsonResponse(
            {
                "lab": lab.id,
                "activitate": activitate.id,
                "month": month_int,
                "year": year_int,
                "links": meta.links if meta else "",
                "livrabil": meta.livrabil if meta else "",
                "comentarii": meta.comentarii if meta else "",
            }
        )

    links = str(payload.get("links") or "").strip()
    livrabil = str(payload.get("livrabil") or "").strip()
    comentarii = str(payload.get("comentarii") or "").strip()

    meta, _created = MonthlyMeta.objects.update_or_create(
        user=request.user,
        lab=lab,
        activitate=activitate,
        month=month_int,
        year=year_int,
        defaults={
            "links": links,
            "livrabil": livrabil,
            "comentarii": comentarii,
        },
    )

    return JsonResponse(
        {
            "status": "ok",
            "id": meta.id,
            "lab": meta.lab_id,
            "activitate": meta.activitate_id,
            "month": meta.month,
            "year": meta.year,
            "links": meta.links,
            "livrabil": meta.livrabil,
            "comentarii": meta.comentarii,
        }
    )

@login_required
def monthly_user_entries(request):
    '''Endpoint to get all work entries for the current user in a given month/year and lab.'''
    month = request.GET.get("month")
    year = request.GET.get("year")

    try:
        month = int(month)
        year = int(year)
    except (TypeError, ValueError):
        return JsonResponse({"error": "Invalid month/year"}, status=400)
    
    lab_id = request.GET.get("lab")
    profile = request.user.userprofile
    can_see_jurnal = LabMembership.objects.filter(
        profile=profile,
        lab__name__in=LAB_JURNAL_NAMES
    ).exists()

    entries = WorkEntry.objects.filter(
        user=request.user,
        lab_id=lab_id,
        date__month=month,
        date__year=year
    ).order_by("-date")

    data = [
        {
            "id": e.id,
            "date": e.date.strftime("%d-%m-%Y"),
            "nr_ore": e.nr_ore,
            "lab": e.lab.name if e.lab else "",
            "activitate": e.activitate.nume if e.activitate else "",
            "livrabil": e.livrabil or "",
            "durata": e.durata,
            "activity_description": e.activity_description,
            "individual": e.individual,
            "members": [u.username for u in e.members.all()],
            "links": e.links,
            "comentarii": e.comentarii,
            **(
                {
                    "jurnal": e.jurnal or "",
                    "scurta_descriere_jurnal": e.scurta_descriere_jurnal or "",
                }
                if (can_see_jurnal)
                else {}
            ),
        }
        for e in entries
    ]

    return JsonResponse(data, safe=False)


@login_required
def generate_jurnal_docx(request):
    """
    Generate a .docx journal for the logged-in user using their own entries.
    Visibility: only users who can see Jurnal for Lab 2.
    """
    lab_id = request.GET.get("lab")
    month = request.GET.get("month")
    year = request.GET.get("year")
    luna = export_views.months_to_RO(int(month))


    try:
        month = int(month)
        year = int(year)
    except (TypeError, ValueError):
        return JsonResponse({"error": "Invalid month/year"}, status=400)

    profile = request.user.userprofile
    can_see_jurnal = LabMembership.objects.filter(
        profile=profile,
        lab__name__in=LAB_JURNAL_NAMES
    ).exists()
    if not can_see_jurnal:
        return HttpResponseForbidden("Nu ai acces la jurnal pentru acest lab.")

    try:
        from docx import Document  # type: ignore
        from docx.oxml import OxmlElement  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
        from docx.opc.constants import RELATIONSHIP_TYPE as RT  # type: ignore
        from docx.shared import Inches, RGBColor  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    except ModuleNotFoundError:
        return JsonResponse(
            {"error": "Missing dependency: python-docx. Install it and retry."},
            status=500,
        )

    def set_document_font_times_new_roman(document):
        style = document.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        # Ensure the font applies to all script types as well.
        rfonts = style.element.rPr.rFonts
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:cs"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "Times New Roman")

    def add_hyperlink(paragraph, url, text):
        # python-docx doesn't expose hyperlinks directly; construct the XML.
        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        run = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")

        # Black, underlined link (Word default look, but explicit).
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "467886")
        r_pr.append(color)

        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)

        run.append(r_pr)
        t = OxmlElement("w:t")
        t.text = text
        run.append(t)
        hyperlink.append(run)

        paragraph._p.append(hyperlink)

    entries = (
        WorkEntry.objects.filter(
            user=request.user,
            lab_id=lab_id,
            date__month=month,
            date__year=year,
        )
        .order_by("date")
    )

    doc = Document()
    set_document_font_times_new_roman(doc)

    for e in entries:
        date_str = f"{e.date.day} {luna} {e.date.year}"
        durata_str = (e.durata or "").strip()

        line2 = (e.scurta_descriere_jurnal or "").strip()
        jurnal_value = (e.jurnal or "").strip()

        p1 = doc.add_paragraph()
        p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(f"{date_str} {durata_str}".strip())
        r1.bold = True
        r1.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

        if line2:
            p2 = doc.add_paragraph()
            p2.add_run(line2)

        p3 = doc.add_paragraph()
        p3.add_run("link to my drive: ")
        if jurnal_value.startswith("http://") or jurnal_value.startswith("https://"):
            add_hyperlink(p3, jurnal_value, jurnal_value)
        else:
            p3.add_run(jurnal_value)

        doc.add_paragraph("")  # spacer

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"jurnal_lab{lab_id}_{year}-{month:02d}_{request.user.username}.docx"
    resp = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp["Content-Disposition"] = f'attachment; filename="{filename}"'
    return resp


@login_required
@require_http_methods(["DELETE", "PATCH"])
def work_entry_detail(request, entry_id):
    """
    DELETE: delete a work entry (owner-only).
    PATCH: update selected fields in-place (owner-only).
    """
    entry = get_object_or_404(WorkEntry, id=entry_id, user=request.user)

    if request.method == "DELETE":
        entry.delete()
        return JsonResponse({"status": "ok"})

    try:
        payload = json.loads(request.body or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"error": "Invalid JSON"}, status=400)

    allowed_fields = {"date", "nr_ore", "durata", "individual", "jurnal", "scurta_descriere_jurnal"}
    updates = {k: payload.get(k) for k in allowed_fields if k in payload}
    if not updates:
        return JsonResponse({"error": "No editable fields provided."}, status=400)

    def _parse_date(value):
        s = str(value or "").strip()
        if not s:
            return None
        # Accept YYYY-MM-DD or DD-MM-YYYY
        try:
            if len(s) >= 10 and s[4] == "-" and s[7] == "-":
                return datetime.strptime(s[:10], "%Y-%m-%d").date()
        except Exception:
            pass
        try:
            return datetime.strptime(s[:10], "%d-%m-%Y").date()
        except Exception:
            return None

    def _duration_hours(value):
        s = str(value or "").strip()
        if "-" not in s:
            return None
        a, b = s.split("-", 1)
        try:
            sh, sm = [int(x) for x in a.strip().split(":", 1)]
            eh, em = [int(x) for x in b.strip().split(":", 1)]
        except Exception:
            return None
        start = sh * 60 + sm
        end = eh * 60 + em
        if end <= start:
            return None
        diff = end - start
        if diff % 60 != 0:
            return None
        return diff // 60

    new_date = entry.date
    if "date" in updates:
        parsed = _parse_date(updates["date"])
        if parsed is None:
            return JsonResponse({"error": "Invalid date format."}, status=400)
        if parsed.weekday() >= 5:
            return JsonResponse({"error": "Ați încercat să pontați in weekend"}, status=400)
        new_date = parsed

    # Validate/update nr_ore with monthly limit checks.
    new_nr_ore = entry.nr_ore
    if "nr_ore" in updates:
        try:
            new_nr_ore = int(updates["nr_ore"])
        except (TypeError, ValueError):
            return JsonResponse({"error": "nr_ore must be an integer"}, status=400)
        if new_nr_ore < 1 or new_nr_ore > 12:
            return JsonResponse({"error": "nr_ore must be between 1 and 12"}, status=400)
        # If nr_ore changes, require durata to be updated/consistent too.
        if "durata" not in updates:
            existing_h = _duration_hours(entry.durata)
            if existing_h != new_nr_ore:
                return JsonResponse(
                    {"error": "Durata nu corespunde cu noul nr_ore. Actualizează și durata."},
                    status=400,
                )

    if "durata" in updates:
        durata = str(updates["durata"] or "").strip()
        if "-" not in durata:
            return JsonResponse({"error": "Durata format invalid."}, status=400)
        h = _duration_hours(durata)
        if h is None:
            return JsonResponse({"error": "Durata format invalid."}, status=400)
        # If either nr_ore changes or durata changes, ensure consistency.
        if "nr_ore" in updates and h != new_nr_ore:
            return JsonResponse({"error": "Durata nu corespunde cu nr_ore."}, status=400)
        if "nr_ore" not in updates and h != entry.nr_ore:
            return JsonResponse({"error": "Durata nu corespunde cu nr_ore."}, status=400)

        entry.durata = durata

    # Apply date + nr_ore and enforce monthly limit for the *target* month/year.
    date_changed = new_date != entry.date
    nr_changed = ("nr_ore" in updates) and (new_nr_ore != entry.nr_ore)
    if date_changed or nr_changed:
        profile = request.user.userprofile
        membership = LabMembership.objects.filter(profile=profile, lab=entry.lab).first()
        limit = membership.monthly_hour_limit if membership is not None else profile.monthly_hour_limit

        total_other = (
            WorkEntry.objects.filter(
                user=request.user,
                lab=entry.lab,
                date__year=new_date.year,
                date__month=new_date.month,
            )
            .exclude(id=entry.id)
            .aggregate(total=Sum("nr_ore"))
            .get("total")
            or 0
        )

        if total_other + new_nr_ore > limit:
            remaining = limit - total_other
            return JsonResponse(
                {"error": f"Ți-ai atins limita lunară. Mai poți introduce doar {remaining} ore."},
                status=400,
            )

        entry.date = new_date
        entry.nr_ore = new_nr_ore

    if "individual" in updates:
        raw = updates["individual"]
        if isinstance(raw, bool):
            entry.individual = raw
        else:
            entry.individual = str(raw).lower() in {"true", "1", "da", "yes"}

    profile = request.user.userprofile
    can_see_jurnal = LabMembership.objects.filter(
        profile=profile,
        lab__name__in=LAB_JURNAL_NAMES
    ).exists()

    if can_see_jurnal:
        if "jurnal" in updates:
            entry.jurnal = str(updates["jurnal"] or "")
        if "scurta_descriere_jurnal" in updates:
            entry.scurta_descriere_jurnal = str(updates["scurta_descriere_jurnal"] or "")

    entry.save()

    resp = {
        "id": entry.id,
        "date": entry.date.strftime("%d-%m-%Y"),
        "nr_ore": entry.nr_ore,
        "durata": entry.durata,
        "lab": entry.lab.name if entry.lab else "",
        "activitate": entry.activitate.nume if entry.activitate else "",
        "activity_description": entry.activity_description,
        "individual": entry.individual,
        "links": entry.links,
        "livrabil": entry.livrabil or "",
        "comentarii": entry.comentarii,
        **(
            {
                "jurnal": entry.jurnal or "",
                "scurta_descriere_jurnal": entry.scurta_descriere_jurnal or "",
            }
            if can_see_jurnal
            else {}
        ),
    }
    return JsonResponse(resp)
