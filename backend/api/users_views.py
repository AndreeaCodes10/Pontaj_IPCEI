'''
Views related to user and lab management, such as listing lab members, adding/removing users from labs, and listing all users.

FUNCTIONS:
----------
lab_members(request, lab_id): 
    Lists members of a lab. Accessible by admin and lab members.
update_monthly_hour_limit(request, lab_id, user_id):
    Updates monthly_hour_limit for a lab member. Accessible by admin and lab directors.
'''
import json
from io import BytesIO
from django.http import JsonResponse, HttpResponseForbidden, HttpResponse
from django.shortcuts import get_object_or_404
from django.views.decorators.http import require_http_methods
from django.contrib.auth.decorators import login_required
from django.utils import timezone
from .models import Lab, User, LabMembership

@login_required
def lab_members(request, lab_id):
    '''Endpoint to list active members of a lab. Only admin or lab members can access this endpoint.'''
    lab = get_object_or_404(Lab, id=lab_id)

    profile = request.user.userprofile

    # check membership or admin
    if profile.role != "admin" and not LabMembership.objects.filter(
        lab=lab, profile=profile
    ).exists():
        return HttpResponseForbidden()

    memberships = LabMembership.objects.filter(lab=lab).select_related("profile__user")

    data = [
        {
            "id": m.profile.user.id,
            "username": m.profile.user.username,
            "role": m.role,
            "monthly_hour_limit": m.monthly_hour_limit,
        }
        for m in memberships
    ]

    return JsonResponse(data, safe=False)

@login_required
@require_http_methods(["PATCH", "POST"])
def update_monthly_hour_limit(request, lab_id, user_id):
    """Update monthly_hour_limit for a user within a lab.

    Permissions:
    - Admins can update any lab membership.
    - Lab directors can update membership limits for their lab.
    """

    lab = get_object_or_404(Lab, id=lab_id)
    target_user = get_object_or_404(User, id=user_id)

    requester_profile = request.user.userprofile

    is_admin = requester_profile.role == "admin"
    is_director_in_lab = LabMembership.objects.filter(
        lab=lab, profile=requester_profile, role="director"
    ).exists()

    if not (is_admin or is_director_in_lab):
        return HttpResponseForbidden()

    try:
        payload = json.loads(request.body.decode("utf-8") or "{}")
    except json.JSONDecodeError:
        return JsonResponse({"status": "error", "message": "Invalid JSON"}, status=400)

    if "monthly_hour_limit" not in payload:
        return JsonResponse(
            {"status": "error", "message": "monthly_hour_limit is required"},
            status=400,
        )

    try:
        new_limit = int(payload["monthly_hour_limit"])
    except (TypeError, ValueError):
        return JsonResponse(
            {"status": "error", "message": "monthly_hour_limit must be an integer"},
            status=400,
        )

    if new_limit < 0 or new_limit > 1000:
        return JsonResponse(
            {"status": "error", "message": "monthly_hour_limit out of range"},
            status=400,
        )

    membership = get_object_or_404(
        LabMembership, lab=lab, profile=target_user.userprofile
    )
    membership.monthly_hour_limit = new_limit
    membership.save(update_fields=["monthly_hour_limit"])

    return JsonResponse({"status": "ok", "monthly_hour_limit": membership.monthly_hour_limit})

def hours_to_hhmm(hours_float):
    hours = int(hours_float)
    minutes = int(round((hours_float - hours) * 60))
    return f"{hours}:{minutes:02d}"

@login_required
@require_http_methods(["GET"])
def generate_anexa1_referat_modificare_docx(request, lab_id, user_id):
    """
    Generate "ANEXA 1 Referat Modificare" .docx for a lab member.
    """

    lab = get_object_or_404(Lab, id=lab_id)
    target_user = get_object_or_404(User, id=user_id)
    requester_profile = request.user.userprofile

    is_admin = requester_profile.role == "admin"
    is_director_in_lab = LabMembership.objects.filter(
        lab=lab, profile=requester_profile, role="director"
    ).exists()

    if not (is_admin or is_director_in_lab):
        return HttpResponseForbidden()

    membership = get_object_or_404(
        LabMembership, lab=lab, profile=target_user.userprofile
    )

    try:
        from docx import Document  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    except ModuleNotFoundError:
        return JsonResponse(
            {"error": "Missing dependency: python-docx. Install it and retry."},
            status=500,
        )
    director_first = (request.user.first_name or "").strip()
    director_last = (request.user.last_name or "").strip()
    director_full_name = (
        (f"{director_first} {director_last}").strip() or request.user.username
    )
    date_str = timezone.localdate().strftime("%d.%m.%Y")

    doc = Document()

    doc.add_paragraph("									                          Anexa 1")
    doc.add_paragraph("Universitatea Politehnica Timișoara")
    doc.add_paragraph("Departamentul EA")
    doc.add_paragraph("")
    doc.add_paragraph("")

    doc.add_paragraph("                APROBAT	                  	     				VIZA*")
    doc.add_paragraph("                  RECTOR		                        Direcția Financiar-Contabilă/Birou Cercetare")
    doc.add_paragraph("Conf.univ.dr.ing. Florin DRĂGAN")
    doc.add_paragraph("")

    doc.add_paragraph("Către").alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Conducerea Universității Politehnica Timișoara").alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    doc.add_paragraph(f"\t {director_full_name}, în calitate de director proiect cu titlul \"{lab.titlu})\", vă rog să aprobați pentru  colaboratorii externi, \
 incluși în tabelul de mai jos, care desfășoară activități în cadrul contractului de cercetare/proiectului menționat, \
 modificarea contractului existent de muncă cu timp parțial de lucru, pe durată determinată.  \n \
 \tMenționez că aceste sume brute au fost calculate pe principiul utilizării sumei reprezentând TOTAL \
 CHELTUIELI DE PERSONAL din devizul contractului și a numărului de ore alocate acestuia.")

    table = doc.add_table(rows=2, cols=9)
    table.style = "Table Grid"

    header = table.rows[0].cells
    body = table.rows[1].cells

    header[0].text = "Nr. crt"
    header[1].text = "Nume Prenume"
    header[2].text = "Funcţia conform statului de funcţii al Contractului"
    header[3].text = "Locul de muncă(al Contractului)"
    header[4].text = "Durata CIM/ actului adiţional(de la…până la…)"
    header[5].text = "Număr de ore prestate pe lună"
    header[6].text = "Durata muncii (ore/zi)"
    header[7].text = "Interval orar"
    header[8].text = "Salariu brut lunar (nr. de ore x tariful orar)"

    body[1].text = target_user.get_full_name().strip()
    body[2].text = (membership.post or "").strip()
    body[3].text = "UPT"
    body[5].text = str(membership.monthly_hour_limit)

    daily_hours = membership.monthly_hour_limit / 20
    body[6].text = hours_to_hhmm(daily_hours)

    start_hour = 16
    end_time = start_hour + daily_hours
    body[7].text = f"{hours_to_hhmm(start_hour)}-{hours_to_hhmm(end_time)}"
    doc.add_paragraph(f"Salar brut = Total cheltuieli de personal /1,0225 \n \
        Menționez că activitățile ce se vor desfășura în cadrul contractului vor fi detaliate în fișa postului, pentru fiecare persoană. \n \
        De asemenea precizez că plata salariilor și a contribuțiilor datorate către Stat atât pentru angajat cât și pentru angajator\
(total cheltuieli de personal) se vor face exclusiv din bugetul contractului de cercetare/finanțare cu titlul \"{lab.titlu}\".")


    doc.add_paragraph(f"     Timișoara,")
    doc.add_paragraph(f"    {date_str}")
    doc.add_paragraph(f"Director proiect").alignment=WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph(f"{director_full_name}").alignment=WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("")
    doc.add_paragraph("         Verificat")
    doc.add_paragraph("Direcția Resurse Umane")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = "ANEXA 1 Referat Modificare.docx"
    resp = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    resp["Content-Disposition"] = f'attachment; filename="{filename}"'
    return resp
